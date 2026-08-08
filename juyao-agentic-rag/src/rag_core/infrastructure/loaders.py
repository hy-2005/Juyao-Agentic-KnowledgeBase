# 原始文档加载：TXT / Markdown / PDF / Word(.docx) / CSV 等 → 统一 UTF-8 纯文本再切块。
# 说明：.md 按文本读入，保留 # 标题与列表符号，便于语义切块；.doc 不支持，请另存为 .docx 或 PDF。

import csv
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def load_text(path: str) -> str:
    # 读取本地文本文件，自动尝试常见编码。
    #
    # 优先顺序：
    # 1) utf-8
    # 2) utf-16（含 BOM 的常见 Windows 文本）
    # 3) gbk（部分中文本地导出文本）
    file_path = Path(path)
    if not file_path.exists():
        raise FileNotFoundError(f"未找到输入文件: {path}")

    tried: list[str] = []
    for encoding in ("utf-8", "utf-16", "gbk"):
        tried.append(encoding)
        try:
            # 成功读取后立即返回，不继续尝试后续编码。
            return file_path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue

    raise UnicodeDecodeError(
        "unknown",
        b"",
        0,
        1,
        f"无法解析文件编码: {path}。已尝试: {', '.join(tried)}。请先转为 UTF-8 再导入。",
    )


def _load_docx_as_text(path: str) -> str:
    try:
        import docx  # python-docx
    except ImportError as exc:
        raise ImportError("读取 Word(.docx) 需要安装 python-docx：pip install python-docx") from exc
    document = docx.Document(path)
    blocks: list[str] = []
    for p in document.paragraphs:
        t = (p.text or "").strip()
        if t:
            blocks.append(t)
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                blocks.append("\t".join(cells))
    return "\n\n".join(blocks).strip() or ""


# OCR 页面文本层判定阈值：低于该字符数且页面含图片视为扫描件页
_OCR_TEXT_THRESHOLD = 20


def _get_ocr_engine():
    """懒加载 RapidOCR 单例；未安装时返回 None 并告警（不阻断入库）。

    扫描件 PDF 无文本层，纯 get_text 会静默丢内容，必须走 OCR。
    """
    if not hasattr(_get_ocr_engine, "_engine"):
        try:
            from rapidocr_onnxruntime import RapidOCR

            _get_ocr_engine._engine = RapidOCR()
            logger.info("RapidOCR 已加载（扫描件 PDF 识别启用）")
        except ImportError:
            logger.warning(
                "未安装 rapidocr-onnxruntime，扫描件 PDF 将跳过 OCR（内容为空）。"
                "安装：pip install rapidocr-onnxruntime"
            )
            _get_ocr_engine._engine = None
    return _get_ocr_engine._engine


def _ocr_page(engine, page) -> str:
    """对单页做 OCR：渲染 PNG → rapidocr 识别 → 按行拼接文本。"""
    try:
        pix = page.get_pixmap(dpi=200)
        result, _ = engine(pix.tobytes("png"))
        if not result:
            return ""
        lines: list[str] = []
        for item in result:
            # rapidocr 返回 [[bbox, text, score], ...]
            if isinstance(item, (list, tuple)) and len(item) >= 2:
                text = str(item[1]).strip()
                if text:
                    lines.append(text)
        return "\n".join(lines)
    except Exception as exc:
        logger.warning("PDF 第 %s 页 OCR 失败：%s", page.number + 1, exc)
        return ""


def _has_scanned_pages(path: str) -> bool:
    """预扫描：是否存在无文本层页面（扫描件页）。全文本层 → 可走布局感知解析。"""
    import fitz

    doc = fitz.open(path)
    try:
        for page in doc:
            if len((page.get_text() or "").strip()) < _OCR_TEXT_THRESHOLD:
                return True
        return False
    finally:
        doc.close()


def _load_pdf_structured(path: str) -> str | None:
    """布局感知 PDF 解析（PyMuPDF4LLM）：表格还原为 Markdown、标题/段落结构保留。

    仅适用于全文本层 PDF（电子版）；失败返回 None 由调用方回退逐页逻辑。
    返回的 Markdown 会再经跨页表格合并（_merge_cross_page_tables）。
    """
    try:
        import pymupdf4llm
    except ImportError:
        logger.info("未安装 pymupdf4llm，PDF 走逐页提取")
        return None
    try:
        md = pymupdf4llm.to_markdown(path, pages=None)
        if not md or not md.strip():
            return None
        merged = _merge_cross_page_tables(md)
        logger.info("PyMuPDF4LLM 布局解析完成：%s 字符（跨页表格合并后）", len(merged))
        return merged
    except Exception as exc:
        logger.warning("PyMuPDF4LLM 解析失败，回退逐页提取：%s", exc)
        return None


def _merge_cross_page_tables(md: str) -> str:
    """跨页表格合并：把分页拆开的表格碎片合并回逻辑整体。

    规则（内容启发式，不依赖几何位置）：
    - 表格块 = 以 | 开头的连续行
    - 相邻表格块之间若被「续表/（续）」标记或重复表头隔开 → 合并：
      后一块首行若与前一表格的表头行相同，视为重复表头，丢弃后再拼接
    """
    lines = md.split("\n")
    merged: list[str] = []
    i = 0
    n = len(lines)

    def is_table_line(s: str) -> bool:
        return s.strip().startswith("|")

    while i < n:
        if not is_table_line(lines[i]):
            merged.append(lines[i])
            i += 1
            continue
        # 收集当前表格块
        block = [lines[i]]
        j = i + 1
        while j < n and is_table_line(lines[j]):
            block.append(lines[j])
            j += 1
        # 向后探测跨页碎片：跳过空行/续标记后若又是表格块且表头相同 → 并入
        k = j
        while k < n:
            line = lines[k].strip()
            if not line:
                k += 1
                continue
            if "续" in line or "continued" in line.lower():
                k += 1
                continue
            if is_table_line(lines[k]) and lines[k].strip() == block[0].strip():
                k += 1
                while k < n and is_table_line(lines[k]):
                    block.append(lines[k])
                    k += 1
            else:
                break
        merged.extend(block)
        if merged and merged[-1]:
            merged.append("")  # 表格块之间补空行,保持 Markdown 结构
        i = k if k > j else j
    return "\n".join(merged)


def load_document(path: str) -> str:
    """按扩展名加载为 UTF-8 纯文本；未知扩展名走 load_text（多编码兜底）。"""
    file_path = Path(path)
    if not file_path.exists():
        raise FileNotFoundError(f"未找到输入文件: {path}")
    suf = file_path.suffix.lower()
    if suf == ".pdf":
        # 全文本层电子版 PDF → 布局感知解析（表格还原为 Markdown，方案 A）
        if not _has_scanned_pages(path):
            structured = _load_pdf_structured(path)
            if structured is not None:
                return structured
        try:
            import fitz  # PyMuPDF
        except ImportError as exc:
            raise ImportError("读取 PDF 需要安装 pymupdf：pip install pymupdf") from exc
        doc = fitz.open(path)
        try:
            parts: list[str] = []
            ocr_engine = None
            for page in doc:
                text = page.get_text() or ""
                # 无文本层的页面（扫描件）→ OCR 兜底，避免内容静默丢失。
                # 注意：整页图片渲染的扫描件 get_images() 可能返回 0（页面级图片不是
                # 嵌入图片对象），不能以"含图片"为触发条件——只要文本过少就尝试 OCR，
                # 纯空白页 OCR 结果为空，不会比不 OCR 更差。
                if len(text.strip()) < _OCR_TEXT_THRESHOLD:
                    if ocr_engine is None:
                        ocr_engine = _get_ocr_engine()
                    if ocr_engine is not None:
                        ocr_text = _ocr_page(ocr_engine, page)
                        if ocr_text:
                            logger.info(
                                "PDF 第 %s 页无文本层，OCR 提取 %s 字符",
                                page.number + 1,
                                len(ocr_text),
                            )
                            text = ocr_text
                parts.append(text)
            return "\n\n".join(parts).strip() or ""
        finally:
            doc.close()
    if suf == ".docx":
        return _load_docx_as_text(path)
    if suf == ".csv":
        lines: list[str] = []
        for enc in ("utf-8-sig", "utf-8", "gbk"):
            try:
                with open(path, newline="", encoding=enc) as f:
                    reader = csv.reader(f)
                    for row in reader:
                        lines.append("\t".join(cell.strip() for cell in row))
                return "\n".join(lines)
            except UnicodeDecodeError:
                continue
        raise UnicodeDecodeError("unknown", b"", 0, 1, f"无法解析 CSV 编码: {path}")
    # 其余（.txt / .md / .markdown / .json / .xml / .html …）：多编码文本读入，Markdown 保留 # 标题等符号
    return load_text(path)
