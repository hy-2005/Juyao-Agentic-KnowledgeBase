# 原始文档加载：TXT / Markdown / PDF / Word(.docx) / CSV 等 → 统一 UTF-8 纯文本再切块。
# 说明：.md 按文本读入，保留 # 标题与列表符号，便于语义切块；.doc 不支持，请另存为 .docx 或 PDF。

import csv
from pathlib import Path


def load_text(path: str) -> str:
    # 读取本地文本文件，自动尝试常见编码。
    #
    # 优先顺序：
    # 1) utf-8
    # 2) utf-16（含 BOM 的常见 Windows 文本）
    # 3) gbk（部分中文本地导出文本）
    file_path = Path(path)
    if not file_path.exists():
        raise FileNotFoundError(f"未找到输入文件: {path}")

    tried: list[str] = []
    for encoding in ("utf-8", "utf-16", "gbk"):
        tried.append(encoding)
        try:
            # 成功读取后立即返回，不继续尝试后续编码。
            return file_path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue

    raise UnicodeDecodeError(
        "unknown",
        b"",
        0,
        1,
        f"无法解析文件编码: {path}。已尝试: {', '.join(tried)}。请先转为 UTF-8 再导入。",
    )


def _load_docx_as_text(path: str) -> str:
    try:
        import docx  # python-docx
    except ImportError as exc:
        raise ImportError("读取 Word(.docx) 需要安装 python-docx：pip install python-docx") from exc
    document = docx.Document(path)
    blocks: list[str] = []
    for p in document.paragraphs:
        t = (p.text or "").strip()
        if t:
            blocks.append(t)
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                blocks.append("\t".join(cells))
    return "\n\n".join(blocks).strip() or ""


def load_document(path: str) -> str:
    """按扩展名加载为 UTF-8 纯文本；未知扩展名走 load_text（多编码兜底）。"""
    file_path = Path(path)
    if not file_path.exists():
        raise FileNotFoundError(f"未找到输入文件: {path}")
    suf = file_path.suffix.lower()
    if suf == ".pdf":
        try:
            import fitz  # PyMuPDF
        except ImportError as exc:
            raise ImportError("读取 PDF 需要安装 pymupdf：pip install pymupdf") from exc
        doc = fitz.open(path)
        try:
            parts: list[str] = []
            for page in doc:
                parts.append(page.get_text() or "")
            return "\n\n".join(parts).strip() or ""
        finally:
            doc.close()
    if suf == ".docx":
        return _load_docx_as_text(path)
    if suf == ".csv":
        lines: list[str] = []
        for enc in ("utf-8-sig", "utf-8", "gbk"):
            try:
                with open(path, newline="", encoding=enc) as f:
                    reader = csv.reader(f)
                    for row in reader:
                        lines.append("\t".join(cell.strip() for cell in row))
                return "\n".join(lines)
            except UnicodeDecodeError:
                continue
        raise UnicodeDecodeError("unknown", b"", 0, 1, f"无法解析 CSV 编码: {path}")
    # 其余（.txt / .md / .markdown / .json / .xml / .html …）：多编码文本读入，Markdown 保留 # 标题等符号
    return load_text(path)
